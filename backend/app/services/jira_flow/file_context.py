"""Procesa archivos adjuntos para Gemini + Jira."""
import io


def mime_for(filename: str, declared_mime: str | None = None) -> str:
    name = (filename or "").lower()
    if declared_mime and declared_mime != "application/octet-stream":
        return declared_mime
    if name.endswith(".png"): return "image/png"
    if name.endswith((".jpg", ".jpeg")): return "image/jpeg"
    if name.endswith(".webp"): return "image/webp"
    if name.endswith(".gif"): return "image/gif"
    if name.endswith(".pdf"): return "application/pdf"
    if name.endswith(".docx"): return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if name.endswith(".xlsx"): return "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    if name.endswith(".txt"): return "text/plain"
    if name.endswith(".md"): return "text/markdown"
    return "application/octet-stream"


def _extract_docx_text(file_bytes: bytes) -> str:
    try:
        from docx import Document
    except ImportError:
        return "[python-docx no instalado]"
    try:
        doc = Document(io.BytesIO(file_bytes))
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Error extrayendo DOCX: {e}]"


def _extract_xlsx_text(file_bytes: bytes, max_rows_per_sheet: int = 100) -> str:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return "[openpyxl no instalado]"
    try:
        wb = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        parts = []
        for sheet_name in wb.sheetnames[:5]:
            ws = wb[sheet_name]
            parts.append(f"\n## Hoja: {sheet_name}")
            for ix, row in enumerate(ws.iter_rows(values_only=True)):
                if ix >= max_rows_per_sheet:
                    parts.append(f"... ({ws.max_row - max_rows_per_sheet} filas más omitidas)")
                    break
                row_text = " | ".join(str(c) if c is not None else "" for c in row)
                if row_text.replace("|", "").strip():
                    parts.append(row_text)
        return "\n".join(parts)
    except Exception as e:
        return f"[Error extrayendo XLSX: {e}]"


def _extract_pdf_text(file_bytes: bytes, max_pages: int = 50) -> str:
    try:
        from pypdf import PdfReader
    except ImportError:
        return ""
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        parts = []
        for ix, page in enumerate(reader.pages[:max_pages]):
            text = page.extract_text() or ""
            if text.strip():
                parts.append(f"[Página {ix+1}]\n{text}")
        return "\n\n".join(parts)
    except Exception:
        return ""


def process_files(files: list[dict]) -> dict:
    """Procesa archivos uploaded.

    Cada item de `files` es un dict con keys: name, mime, bytes.
    """
    images: list[tuple[bytes, str]] = []
    pdfs: list[tuple[bytes, str, str]] = []
    texts: list[tuple[str, str]] = []
    all_attach: list[tuple[bytes, str, str]] = []

    for f in files or []:
        data = f["bytes"]
        name = f["name"]
        mime = mime_for(name, f.get("mime"))
        all_attach.append((data, mime, name))

        if mime.startswith("image/"):
            images.append((data, mime))
        elif mime == "application/pdf":
            pdfs.append((data, mime, name))
        elif mime.endswith("wordprocessingml.document"):
            texts.append((name, _extract_docx_text(data)))
        elif mime.endswith("spreadsheetml.sheet"):
            texts.append((name, _extract_xlsx_text(data)))
        elif mime in ("text/plain", "text/markdown"):
            try:
                texts.append((name, data.decode("utf-8", errors="replace")))
            except Exception:
                pass

    return {
        "images_for_gemini": images,
        "pdfs_for_gemini": pdfs,
        "extracted_texts": texts,
        "all_attachments": all_attach,
    }


def build_context_with_files(base_context: str, processed: dict) -> str:
    parts = [base_context]
    texts = processed.get("extracted_texts", [])
    if texts:
        parts.append("\n\n---\nARCHIVOS DE CONTEXTO ADJUNTOS (texto extraído):\n")
        for name, txt in texts:
            parts.append(f"\n### 📎 {name}")
            parts.append(txt[:8000])
            parts.append("")
    return "\n".join(parts)
